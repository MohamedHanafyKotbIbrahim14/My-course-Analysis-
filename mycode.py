import streamlit as st
import pandas as pd
import numpy as np
from datetime import datetime
import io
import gdown
import os
import tempfile

st.set_page_config(page_title="Multi-Course Comparison Report", layout="wide")

st.title("📑 Multi-Course Comparison Report")

st.markdown("---")

# Google Drive configuration
# You'll paste your Google Drive folder link here
GOOGLE_DRIVE_FOLDER = st.secrets.get("GOOGLE_DRIVE_FOLDER", "")

# Cache the data loading to avoid repeated downloads
@st.cache_data(ttl=3600)  # Cache for 1 hour
def load_course_data_from_drive(folder_url):
    """
    Load CSV files from Google Drive folder and extract summary statistics
    """
    if not folder_url:
        return []
    
    try:
        # Extract folder ID from URL
        # URL format: https://drive.google.com/drive/folders/FOLDER_ID?usp=sharing
        folder_id = folder_url.split('/folders/')[1].split('?')[0]
        
        # Create temporary directory
        temp_dir = tempfile.mkdtemp()
        
        # Download folder contents
        gdown.download_folder(id=folder_id, output=temp_dir, quiet=False, use_cookies=False)
        
        # Process each CSV file
        courses_data = []
        for filename in os.listdir(temp_dir):
            if filename.endswith('.csv') or filename.endswith('.CSV'):
                file_path = os.path.join(temp_dir, filename)
                try:
                    # Read CSV
                    df = pd.read_csv(file_path)
                    
                    # Process dataframe
                    df = process_dataframe(df)
                    
                    # Extract course name from filename
                    course_name = filename.replace('.csv', '').replace('.CSV', '')
                    
                    # Get statistics
                    stats = get_course_statistics(df, course_name)
                    courses_data.append(stats)
                    
                except Exception as e:
                    st.warning(f"Could not process {filename}: {str(e)}")
        
        # Clean up temp directory
        import shutil
        shutil.rmtree(temp_dir)
        
        return courses_data
        
    except Exception as e:
        st.error(f"Error loading data from Google Drive: {str(e)}")
        return []

def process_dataframe(df):
    """Process and clean the dataframe"""
    # Skip first row if it's header info
    if 'Student ID' in df.columns and len(df) > 0 and str(df.iloc[0, 0]).startswith('ACTL'):
        df = df.iloc[1:].reset_index(drop=True)
    
    # Clean column names
    df.columns = df.columns.str.strip()
    
    # COLUMN D (index 3): ALWAYS contains Final Mark + Grade
    final_col = df.iloc[:, 3]  # Column D
    df['Final_Mark'] = final_col.astype(str).str.extract(r'(\d+)').astype(float)
    df['Grade'] = final_col.astype(str).str.extract(r'([A-Z]{2,3})')[0]
    
    # Clean Student ID
    df['Student ID'] = df.iloc[:, 0].astype(str).str.strip()
    
    return df

def get_grade_distribution(df):
    """Calculate grade distribution percentages"""
    grade_counts = df['Grade'].value_counts()
    total = len(df[df['Grade'].notna()])
    
    grade_order = ['HD', 'DN', 'CR', 'PS', 'FL']
    distribution = {}
    
    for grade in grade_order:
        count = grade_counts.get(grade, 0)
        percentage = (count / total * 100) if total > 0 else 0
        distribution[grade] = {'count': count, 'percentage': percentage}
    
    return distribution

def get_course_statistics(df, course_name):
    """Get comprehensive statistics for a course"""
    stats = df['Final_Mark'].describe()
    grade_dist = get_grade_distribution(df)
    
    return {
        'course': course_name,
        'count': int(stats['count']),
        'mean': stats['mean'],
        'median': stats['50%'],
        'std': stats['std'],
        'min': stats['min'],
        'max': stats['max'],
        'q25': stats['25%'],
        'q75': stats['75%'],
        'hd_pct': grade_dist['HD']['percentage'],
        'dn_pct': grade_dist['DN']['percentage'],
        'cr_pct': grade_dist['CR']['percentage'],
        'ps_pct': grade_dist['PS']['percentage'],
        'fl_pct': grade_dist['FL']['percentage']
    }

# Section 1: LIC Information
st.header("👤 Lecturer in Charge Information")

col1, col2, col3 = st.columns(3)
with col1:
    lic_name = st.text_input("LIC Name", value="Hanafy")
with col2:
    lic_phone = st.text_input("Phone Number", value="+61404488448")
with col3:
    lic_course = st.text_input("Course Code", value="ACTL5115")

st.markdown("---")

# Load data from Google Drive
st.header("📊 Course Data from Google Drive")

if not GOOGLE_DRIVE_FOLDER:
    st.error("⚠️ Google Drive folder not configured!")
    st.info("""
    **To enable automatic data loading:**
    1. Add your Google Drive folder link to Streamlit Secrets
    2. Key: `GOOGLE_DRIVE_FOLDER`
    3. Value: Your folder sharing link
    
    For now, you can test locally by setting the folder link in the code.
    """)
    st.stop()

with st.spinner("Loading course data from Google Drive..."):
    courses_data = load_course_data_from_drive(GOOGLE_DRIVE_FOLDER)

if not courses_data:
    st.warning("No course data found in Google Drive folder.")
    st.info("Make sure your CSV files are uploaded to the shared folder.")
    st.stop()

st.success(f"✅ Loaded {len(courses_data)} courses from Google Drive!")

# Show available courses
st.subheader("📁 Available Courses")
for course in courses_data:
    st.write(f"- **{course['course']}** ({course['count']} students)")

st.markdown("---")

# Generate Report Button
if st.button("📊 Generate Comparison Report", type="primary"):
    # Store in session state
    st.session_state.courses_data = courses_data
    st.session_state.lic_info = {
        'name': lic_name,
        'phone': lic_phone,
        'course_code': lic_course
    }
    st.success(f"✅ Report generated with {len(courses_data)} course(s)!")

# Display comparison table if data exists
if 'courses_data' in st.session_state and st.session_state.courses_data:
    st.markdown("---")
    st.header("📊 Course Comparison Summary")
    
    # Create comparison dataframe
    comparison_data = []
    
    # Statistics rows
    stats_rows = [
        ('Number of Students', 'count', '.0f'),
        ('Mean', 'mean', '.2f'),
        ('Median', 'median', '.2f'),
        ('Std Dev', 'std', '.2f'),
        ('Min', 'min', '.2f'),
        ('Max', 'max', '.2f'),
        ('25th Percentile', 'q25', '.2f'),
        ('75th Percentile', 'q75', '.2f'),
        ('% HD', 'hd_pct', '.1f'),
        ('% DN', 'dn_pct', '.1f'),
        ('% CR', 'cr_pct', '.1f'),
        ('% PS', 'ps_pct', '.1f'),
        ('% FL', 'fl_pct', '.1f')
    ]
    
    for label, key, fmt in stats_rows:
        row = {'Statistic': label}
        for course in st.session_state.courses_data:
            value = course[key]
            if 'pct' in key:
                row[course['course']] = f"{value:{fmt}}%"
            else:
                row[course['course']] = f"{value:{fmt}}"
        # Add empty "This Year" column
        row['This Year'] = ''
        comparison_data.append(row)
    
    comparison_df = pd.DataFrame(comparison_data)
    
    # Display table
    st.dataframe(comparison_df, use_container_width=True, hide_index=True)
    
    # Download options
    st.markdown("---")
    st.header("📥 Download Report")
    
    col_download1, col_download2 = st.columns(2)
    
    with col_download1:
        # CSV Download
        csv = comparison_df.to_csv(index=False)
        st.download_button(
            "⬇️ Download as CSV",
            csv,
            "course_comparison.csv",
            "text/csv",
            use_container_width=True
        )
    
    with col_download2:
        # Word Document Download using Python
        if st.button("📄 Generate Word Document", use_container_width=True):
            with st.spinner("Generating Word document..."):
                try:
                    from docx import Document
                    from docx.shared import Inches, Pt, RGBColor
                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                    from docx.oxml.ns import qn
                    from docx.oxml import OxmlElement
                    
                    # Create document
                    doc = Document()
                    
                    # Set default font
                    style = doc.styles['Normal']
                    font = style.font
                    font.name = 'Arial'
                    font.size = Pt(11)
                    
                    # Title
                    title = doc.add_paragraph()
                    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    title_run = title.add_run("Course Comparison Report")
                    title_run.font.size = Pt(28)
                    title_run.font.bold = True
                    title_run.font.color.rgb = RGBColor(0, 0, 0)
                    
                    doc.add_paragraph()  # Spacing
                    
                    # LIC Information Section
                    heading = doc.add_paragraph()
                    heading_run = heading.add_run("Lecturer in Charge Information")
                    heading_run.font.size = Pt(16)
                    heading_run.font.bold = True
                    heading_run.font.color.rgb = RGBColor(0, 0, 0)
                    
                    # LIC details
                    p1 = doc.add_paragraph()
                    p1.add_run("Name: ").bold = True
                    p1.add_run(st.session_state.lic_info['name'])
                    
                    p2 = doc.add_paragraph()
                    p2.add_run("Phone: ").bold = True
                    p2.add_run(st.session_state.lic_info['phone'])
                    
                    p3 = doc.add_paragraph()
                    p3.add_run("Course Code: ").bold = True
                    p3.add_run(st.session_state.lic_info['course_code'])
                    
                    p4 = doc.add_paragraph()
                    p4.add_run("Report Generated: ").bold = True
                    p4.add_run(datetime.now().strftime('%B %d, %Y at %H:%M'))
                    
                    doc.add_paragraph()  # Spacing
                    
                    # Comparison Table Section
                    heading2 = doc.add_paragraph()
                    heading2_run = heading2.add_run("Course Comparison")
                    heading2_run.font.size = Pt(16)
                    heading2_run.font.bold = True
                    heading2_run.font.color.rgb = RGBColor(0, 0, 0)
                    
                    # Create table
                    num_courses = len(st.session_state.courses_data)
                    table = doc.add_table(rows=14, cols=num_courses + 2)  # +2 for Statistic column and This Year column
                    table.style = 'Light Grid Accent 1'
                    
                    # Helper function to shade cells
                    def shade_cell(cell, color):
                        shading_elm = OxmlElement('w:shd')
                        shading_elm.set(qn('w:fill'), color)
                        cell._element.get_or_add_tcPr().append(shading_elm)
                    
                    # Header row
                    header_cells = table.rows[0].cells
                    header_cells[0].text = "Statistic"
                    for idx, course in enumerate(st.session_state.courses_data):
                        header_cells[idx + 1].text = course['course']
                    # Add "This Year" as last column
                    header_cells[num_courses + 1].text = "This Year"
                    
                    # Style header row
                    for cell in header_cells:
                        shade_cell(cell, "4472C4")
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                                run.font.color.rgb = RGBColor(255, 255, 255)
                                run.font.size = Pt(11)
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Data rows
                    stats_data = [
                        ('Number of Students', 'count', lambda v: f'{int(v)}'),
                        ('Mean', 'mean', lambda v: f'{v:.2f}'),
                        ('Median', 'median', lambda v: f'{v:.2f}'),
                        ('Std Dev', 'std', lambda v: f'{v:.2f}'),
                        ('Min', 'min', lambda v: f'{v:.2f}'),
                        ('Max', 'max', lambda v: f'{v:.2f}'),
                        ('25th Percentile', 'q25', lambda v: f'{v:.2f}'),
                        ('75th Percentile', 'q75', lambda v: f'{v:.2f}'),
                        ('% HD', 'hd_pct', lambda v: f'{v:.1f}%'),
                        ('% DN', 'dn_pct', lambda v: f'{v:.1f}%'),
                        ('% CR', 'cr_pct', lambda v: f'{v:.1f}%'),
                        ('% PS', 'ps_pct', lambda v: f'{v:.1f}%'),
                        ('% FL', 'fl_pct', lambda v: f'{v:.1f}%')
                    ]
                    
                    for row_idx, (label, key, formatter) in enumerate(stats_data, start=1):
                        row_cells = table.rows[row_idx].cells
                        
                        # First column - statistic name
                        row_cells[0].text = label
                        row_cells[0].paragraphs[0].runs[0].font.bold = True
                        
                        # Data columns
                        for col_idx, course in enumerate(st.session_state.courses_data):
                            row_cells[col_idx + 1].text = formatter(course[key])
                            row_cells[col_idx + 1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Add empty "This Year" cell
                        row_cells[num_courses + 1].text = ""
                        row_cells[num_courses + 1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Alternate row shading
                        if row_idx % 2 == 0:
                            for cell in row_cells:
                                shade_cell(cell, "F2F2F2")
                    
                    # Save to bytes
                    doc_io = io.BytesIO()
                    doc.save(doc_io)
                    doc_io.seek(0)
                    
                    st.success("✅ Word document generated successfully!")
                    st.download_button(
                        "⬇️ Download Word Document",
                        doc_io.getvalue(),
                        "course_comparison_report.docx",
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                
                except ImportError:
                    st.error("❌ python-docx package not installed")
                    st.markdown("""
                    **To enable Word export, install python-docx:**
                    ```
                    pip install python-docx
                    ```
                    
                    Then refresh this page and try again.
                    """)
                except Exception as e:
                    st.error(f"Error generating document: {str(e)}")
                    import traceback
                    st.code(traceback.format_exc())

else:
    st.info("👆 Click 'Generate Comparison Report' to see the results!")
